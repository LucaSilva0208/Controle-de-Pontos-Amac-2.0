import calendar
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import holidays
from gerador_folha import GeradorBase, MESES, DIAS_SEMANA

class GeradorFolha12x36(GeradorBase):
    def gerar(self, funcionario, mes, ano, caminho_saida_pdf, recessos_globais=None):
        self.validar_regras_negocio(mes, ano)
        feriados_br = holidays.BR(subdiv='MG', years=ano)
        data_corpus = self.calcular_corpus_christi(ano)
        feriados_br.append({data_corpus: "Corpus Christi"})
        if recessos_globais is None: recessos_globais = []

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(0.5); section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.0); section.right_margin = Cm(1.0)

        mes_extenso = MESES[int(mes)]
        self._inserir_cabecalho(doc)
        self.add_spacer(doc, pt_size=1)

        # DADOS (Layout 12x36 - Similar ao 40h mas com escala explícita)
        table_info = doc.add_table(rows=3, cols=2)
        table_info.style = 'Table Grid'
        unidade = funcionario.get('unidade', funcionario.get('lotacao', ''))

        self.celula_texto(table_info.cell(0, 0), f"Nome: {funcionario['nome']}", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        self.celula_texto(table_info.cell(0, 1), f"Matrícula: {funcionario['matricula']}", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        # Mescla a célula de lotação para ocupar a linha toda
        table_info.cell(1, 0).merge(table_info.cell(1, 1))
        self.celula_texto(table_info.cell(1, 0), f"Lotação: {unidade}", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        self.celula_texto(table_info.cell(2, 0), f"Cargo: {funcionario.get('cargo', '')}", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        
        self.add_spacer(doc, pt_size=2)

        # --- NOVA TABELA DE HORÁRIOS E DATA ---
        table_horario = doc.add_table(rows=2, cols=6)
        table_horario.style = 'Table Grid'
        
        # Larguras ajustadas para totalizar 19.8cm
        larguras_h = [2.65, 2.65, 2.65, 2.65, 5.3, 3.7]
        for row in table_horario.rows:
            for i, w in enumerate(larguras_h):
                row.cells[i].width = Cm(w)
        
        # Linha 1: Cabeçalhos
        r1 = table_horario.rows[0]
        r1.cells[0].merge(r1.cells[1]); self.celula_texto(r1.cells[0], "Primeira Jornada", bold=True, size=8)
        r1.cells[2].merge(r1.cells[3]); self.celula_texto(r1.cells[2], "Segunda Jornada", bold=True, size=8)
        self.celula_texto(r1.cells[4], "Mês de Referência", bold=True, size=8)
        self.celula_texto(r1.cells[5], "Ano", bold=True, size=8)

        # Linha 2: Dados Puros
        r2 = table_horario.rows[1]
        
        def fmt_hora(h):
            s = str(h) if h else ""
            # Se o texto for longo (ex: 07:00:00), pega apenas os 5 primeiros caracteres (07:00)
            return s[:5] if len(s) >= 8 and ":" in s else s

        vals = [fmt_hora(funcionario.get('horario_ent1','')), fmt_hora(funcionario.get('horario_sai1','')), fmt_hora(funcionario.get('horario_ent2','')), fmt_hora(funcionario.get('horario_sai2','')), mes_extenso, str(ano)]
        for i, val in enumerate(vals): self.celula_texto(r2.cells[i], val, size=9)

        self.add_spacer(doc, pt_size=2)

        # TABELA 12x36 (12 Colunas)
        table = doc.add_table(rows=0, cols=12)
        table.style = 'Table Grid'
        table.autofit = False; table.allow_autofit = False
        # Larguras padrão (Igual 40h)
        larguras = [1.55] * 12

        row_h1 = table.add_row(); self.definir_altura_linha(row_h1, 0.7)
        self.celula_texto(row_h1.cells[0], "Dados", size=7)
        self.celula_texto(row_h1.cells[1], "Primeira Jornada", size=7); row_h1.cells[1].merge(row_h1.cells[4])
        self.celula_texto(row_h1.cells[5], "Segunda Jornada", size=7); row_h1.cells[5].merge(row_h1.cells[8])
        self.celula_texto(row_h1.cells[9], "H. Extra", size=7); row_h1.cells[9].merge(row_h1.cells[10])
        self.celula_texto(row_h1.cells[11], "Visto/Obs", size=7)

        row_h2 = table.add_row(); self.definir_altura_linha(row_h2, 0.5)
        cols_titulos = ["Dia/Sem", "Entr", "Rub", "Saída", "Rub", "Entr", "Rub", "Saída", "Rub", "Entr", "Saída", "Assinatura"]
        for i, tit in enumerate(cols_titulos):
            self.celula_texto(row_h2.cells[i], tit, bold=True, size=6)
            row_h2.cells[i].width = Cm(larguras[i])

        # 12x36 não obedece a recesso global, passamos lista vazia [] para garantir
        self._preencher_dias(table, larguras, mes, ano, feriados_br, [], funcionario)

        self.add_spacer(doc, pt_size=5)
        p = doc.add_paragraph("Reconheço a exatidão destas anotações:")
        p.style.font.size = Pt(8); p.paragraph_format.space_after = Pt(9)
        table_ass = doc.add_table(rows=1, cols=3)
        self.celula_texto(table_ass.cell(0, 0), "_____________________________\nAssinatura Empregado", size=8)
        self.celula_texto(table_ass.cell(0, 1), "Data: ______/______/______", size=8)
        self.celula_texto(table_ass.cell(0, 2), "_____________________________\nAssinatura Chefia", size=8)

        self._salvar_pdf(doc, caminho_saida_pdf)

    def _preencher_dias(self, table, larguras, mes, ano, feriados_br, recessos_globais, funcionario):
        _, num_dias = calendar.monthrange(ano, mes)
        lista_ferias = funcionario.get('ferias', [])

        def verificar_ferias(data_obj, lista_f):
            if not lista_f: return None
            for p in lista_f:
                try:
                    ini = datetime.strptime(p['inicio'], "%Y-%m-%d").date()
                    fim = datetime.strptime(p['fim'], "%Y-%m-%d").date()
                    if ini <= data_obj <= fim:
                        return p.get('tipo', 'FÉRIAS').upper()
                except: continue
            return None

        for dia in range(1, num_dias + 1):
            data = date(ano, mes, dia)
            idx_sem = data.weekday()
            nome_sem = DIAS_SEMANA[idx_sem]
            motivo_afastamento = verificar_ferias(data, lista_ferias)

            row = table.add_row()
            self.definir_altura_linha(row, 0.61) 
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)

            texto_dia_sem = f"{dia:02d} - {nome_sem}"
            self.celula_texto(row.cells[0], texto_dia_sem, size=8)

            bloquear = False
            texto_bloqueio = ""

            if motivo_afastamento:
                bloquear = True
                texto_bloqueio = motivo_afastamento
            # 12x36 NÃO bloqueia FDS, Feriado ou Recesso

            if bloquear:
                for cell in row.cells:
                    self.aplicar_fundo_cinza(cell)
                cell_inicio = row.cells[1]
                cell_fim = row.cells[-1]
                cell_inicio.merge(cell_fim)
                self.celula_texto(cell_inicio, texto_bloqueio, bold=True, size=7)
