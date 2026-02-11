import sys
import os
from docx import Document
from docx.shared import Cm, Pt

def pixels_to_cm(emu):
    """Converte a unidade interna do Word (EMU) para Centímetros"""
    if emu is None: return 0.0
    return round(emu / 360000, 2)

def escanear_modelo(caminho_arquivo):
    if not os.path.exists(caminho_arquivo):
        print(f"ERRO: Arquivo '{caminho_arquivo}' não encontrado.")
        print(f"O script está procurando na pasta: {os.getcwd()}")
        return

    print(f"--- INICIANDO SCAN DO ARQUIVO: {caminho_arquivo} ---\n")
    
    doc = Document(caminho_arquivo)

    # 1. Margens da Página
    section = doc.sections[0]
    print(f"### CONFIGURAÇÃO DA PÁGINA ###")
    print(f"Margem Superior: {pixels_to_cm(section.top_margin)} cm")
    print(f"Margem Inferior: {pixels_to_cm(section.bottom_margin)} cm")
    print(f"Margem Esquerda: {pixels_to_cm(section.left_margin)} cm")
    print(f"Margem Direita:  {pixels_to_cm(section.right_margin)} cm")
    print("-" * 30)

    # 2. Tabelas
    for i, table in enumerate(doc.tables):
        print(f"\n### TABELA {i+1} ###")
        
        # Analisa a primeira linha para pegar larguras das colunas
        if len(table.rows) > 0:
            row = table.rows[0]
            larguras = []
            textos = []
            
            for cell in row.cells:
                larguras.append(pixels_to_cm(cell.width))
                textos.append(cell.text.strip().replace('\n', ' '))
            
            print(f"Total de Colunas: {len(row.cells)}")
            print(f"Larguras (cm): {larguras}")
            print(f"Cabeçalho: {textos}")
            
            # Gera sugestão de código Python
            print("\n>>> SUGESTÃO DE CÓDIGO PARA O SEU SISTEMA:")
            print(f"# Tabela {i+1}")
            print(f"table = doc.add_table(rows=0, cols={len(row.cells)})")
            print(f"larguras = {larguras}")
            print("for i, largura in enumerate(larguras):")
            print("    # Aplique a largura na célula ou coluna")
            print("    pass")
        else:
            print("Tabela vazia.")
        print("-" * 30)

    print("\n--- FIM DO SCAN ---")

if __name__ == "__main__":
    # Garante que procura na mesma pasta onde este script está salvo
    pasta_script = os.path.dirname(os.path.abspath(__file__))
    
    # Define a subpasta dedicada para os arquivos (Agora aponta para 'modelos')
    nome_subpasta = "modelos"
    pasta_scan = os.path.join(pasta_script, nome_subpasta)

    # 1. Cria a pasta se não existir
    if not os.path.exists(pasta_scan):
        print(f"Pasta '{nome_subpasta}' não encontrada em: {pasta_scan}")
        sys.exit()

    # 2. Busca qualquer arquivo .docx dentro dessa pasta
    arquivos_encontrados = [f for f in os.listdir(pasta_scan) if f.lower().endswith(".docx") and not f.startswith("~$")]

    if not arquivos_encontrados:
        print(f"\n>>> NENHUM ARQUIVO ENCONTRADO <<<")
        print(f"A pasta '{nome_subpasta}' não contém arquivos .docx.")
        sys.exit()

    # 3. Prioriza 'modelo_30h.docx' se existir, senão pega o primeiro
    nome_arquivo = "modelo_30h.docx" if "modelo_30h.docx" in arquivos_encontrados else arquivos_encontrados[0]
    arquivo_alvo = os.path.join(pasta_scan, nome_arquivo)
    
    print(f"Arquivo selecionado automaticamente: {nome_arquivo}")
    escanear_modelo(arquivo_alvo)
