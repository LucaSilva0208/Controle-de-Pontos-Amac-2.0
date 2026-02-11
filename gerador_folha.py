import os
import calendar
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx2pdf import convert
import holidays

MESES = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

DIAS_SEMANA = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]

class GeradorBase:
    """Classe base com funcionalidades comuns a todas as folhas"""
    def __init__(self):
        pasta_script = os.path.dirname(os.path.abspath(__file__))
        self.pasta_modelos = os.path.join(pasta_script, "modelos")
        self.logo_path = os.path.join(self.pasta_modelos, "logo.png")
        self.cabecalho_img_path = os.path.join(self.pasta_modelos, "cabecalho.png")

    def validar_regras_negocio(self, mes_alvo, ano_alvo):
        hoje = date.today()
        data_alvo = date(ano_alvo, mes_alvo, 1)
        data_atual_base = date(hoje.year, hoje.month, 1)

        diferenca_meses = (data_atual_base.year - data_alvo.year) * 12 + (data_atual_base.month - data_alvo.month)

        if diferenca_meses > 3:
            raise ValueError(f"Bloqueado: Não é permitido gerar folhas com mais de 3 meses de retrocesso.\nMês solicitado: {mes_alvo}/{ano_alvo}")

        if data_alvo > data_atual_base:
            if diferenca_meses == -1:
                if hoje.day < 20:
                    raise ValueError(f"Bloqueado: A folha do próximo mês ({mes_alvo}/{ano_alvo}) só pode ser gerada a partir do dia 20.")
            elif diferenca_meses < -1:
                 raise ValueError(f"Bloqueado: Não é permitido gerar folhas tão adiante no tempo.")

    def calcular_corpus_christi(self, ano):
        """Calcula Corpus Christi matematicamente (Páscoa + 60 dias)"""
        a = ano % 19
        b = ano // 100
        c = ano % 100
        d = b // 4
        e = b % 4
        f = (b + 8) // 25
        g = (b - f + 1) // 3
        h = (19 * a + b - d - g + 15) % 30
        i = c // 4
        k = c % 4
        l = (32 + 2 * e + 2 * i - h - k) % 7
        m = (a + 11 * h + 22 * l) // 451
        mes_pascoa = (h + l - 7 * m + 114) // 31
        dia_pascoa = ((h + l - 7 * m + 114) % 31) + 1
        
        data_pascoa = date(ano, mes_pascoa, dia_pascoa)
        return data_pascoa + timedelta(days=60)

    def obter_nome_feriado(self, data_obj, feriados_lib):
        fixos = {
            (1, 1): "Confraternização Universal", 
            (21, 4): "Tiradentes", 
            (1, 5): "Dia do Trabalho",
            (13, 6): "Santo Antônio", # Padroeiro JF
            (7, 9): "Independência do Brasil", 
            (12, 10): "Nossa Sra. Aparecida", 
            (2, 11): "Finados",
            (15, 11): "Proclamação da República", 
            (20, 11): "Consciência Negra", 
            (25, 12): "Natal"
        }
        if (data_obj.day, data_obj.month) in fixos: 
            return fixos[(data_obj.day, data_obj.month)]

        nome = feriados_lib.get(data_obj)
        
        if nome:
            traducoes = {
                "Mardi Gras": "Carnaval", "Carnival": "Carnaval", 
                "Good Friday": "Sexta-feira Santa",
                "Corpus Christi": "Corpus Christi", 
                "Easter Sunday": "Páscoa", "Ash Wednesday": "Quarta-feira de Cinzas"
            }
            for eng, pt in traducoes.items():
                if eng in nome:
                    return pt
            return nome
            
        return None

    def aplicar_fundo_cinza(self, cell):
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _salvar_pdf(self, doc, caminho_saida_pdf):
        # --- SALVAR ---
        docx_temp = caminho_saida_pdf.replace(".pdf", ".docx")
        
        dir_saida = os.path.dirname(caminho_saida_pdf)
        if dir_saida:
            os.makedirs(dir_saida, exist_ok=True)
        
        doc.save(docx_temp)
        try:
            convert(docx_temp, caminho_saida_pdf)
        except Exception as e:
            print(f"Erro PDF: {e}")
            raise e
        finally:
            if os.path.exists(docx_temp):
                os.remove(docx_temp)

    # --- MÉTODOS AUXILIARES ---
    def definir_altura_linha(self, row, cm_val):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(cm_val)

    def add_spacer(self, doc, pt_size=5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run()
        run.font.size = Pt(pt_size)

    def _inserir_cabecalho(self, doc):
        """
        Estratégia Simplificada:
        1. Tenta inserir 'cabecalho.png' como uma imagem única e completa (Logo + Texto).
        2. Se a imagem não existir, cria um cabeçalho manual apenas com texto (Fallback).
        """
        if os.path.exists(self.cabecalho_img_path):
            try:
                doc.add_picture(self.cabecalho_img_path, width=Cm(19.9))
                # Ajuste: Centralizar e remover espaçamento padrão do parágrafo da imagem
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_p.paragraph_format.left_indent = Cm(-0.4)
                last_p.paragraph_format.space_after = Pt(0)
                return
            except Exception as e:
                print(f"Erro ao inserir imagem de cabeçalho: {e}")
        
        # Fallback para o manual se a imagem falhar ou não existir
        self._criar_cabecalho_manual(doc)

    def _criar_cabecalho_manual(self, doc):
        """
        Fallback: Cria um cabeçalho simples (apenas texto e logo se houver)
        caso a imagem principal não exista.
        """
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        
        # Coluna da Logo (menor) e Texto (maior)
        header_table.columns[0].width = Cm(3.0)
        header_table.columns[1].width = Cm(19.0) 
        
        cell_logo = header_table.cell(0, 0)
        cell_texto = header_table.cell(0, 1)

        # --- LOGO (Alinhada à Direita) ---
        if os.path.exists(self.logo_path):
            try:
                p = cell_logo.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(self.logo_path, width=Cm(2.5))
            except Exception as e:
                print(f"Erro logo: {e}")
                cell_logo.text = "LOGO"
        else:
            cell_logo.text = "LOGO"
        
        # Centraliza a logo verticalmente
        cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # --- TEXTO CENTRALIZADO COM HIERARQUIA ---
        p_empresa = cell_texto.paragraphs[0]
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_empresa.paragraph_format.space_after = Pt(0)
        
        # 1. Nome da Associação (GRANDE e NEGRITO)
        run_title = p_empresa.add_run("Associação Municipal de Apoio Comunitário\n")
        run_title.bold = True
        run_title.font.size = Pt(14)
        
        # 2. Endereço Completo (PEQUENO e NORMAL)
        run_end = p_empresa.add_run("Rua Espírito Santo, 434 – Centro / Juiz de Fora – MG / Cep: 36010-040\n\n")
        run_end.bold = False
        run_end.font.size = Pt(9)
        
        # 3. Título (GRANDE e NEGRITO)
        run_folha = p_empresa.add_run("FOLHA DE PONTO")
        run_folha.bold = True
        run_folha.font.size = Pt(14)
        
        # Centraliza o texto verticalmente
        cell_texto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def celula_texto(self, cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = text
        if cell.paragraphs:
            p = cell.paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if p.runs:
                run = p.runs[0]
                run.font.bold = bold
                run.font.size = Pt(size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


class GeradorFolhaPonto:
    """Classe Fachada (Facade) para manter compatibilidade com o restante do sistema"""
    def validar_regras_negocio(self, mes, ano):
        GeradorBase().validar_regras_negocio(mes, ano)

    def gerar(self, funcionario, mes, ano, caminho_saida_pdf, recessos_globais=None):
        # Importação local para evitar ciclo, já que as classes filhas importam GeradorBase deste arquivo
        from gerador_folha_30h import GeradorFolha30h
        from gerador_folha_40h import GeradorFolha40h
        from gerador_folha_12x36 import GeradorFolha12x36

        escala = funcionario.get('escala', 'NORMAL')
        carga = funcionario.get('carga_horaria', '40h')
        
        if escala == '12X36':
            gerador = GeradorFolha12x36()
        elif carga == '30h':
            gerador = GeradorFolha30h()
        else:
            gerador = GeradorFolha40h()
        
        gerador.gerar(funcionario, mes, ano, caminho_saida_pdf, recessos_globais)